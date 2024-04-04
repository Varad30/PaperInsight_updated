from django.shortcuts import render
import re
import os
import json
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse, JsonResponse, FileResponse
import google.generativeai as genai
import re
from tika import parser
import docx
import pathlib
from docx import Document
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from pathlib import Path
import matplotlib.pyplot as plt
import csv
import base64
from io import StringIO, BytesIO
from django.shortcuts import render
from io import StringIO
import random
import nltk

nltk.download('stopwords')
nltk.download('punkt')

import convertapi

convertapi.api_secret = ''
genai.configure(api_key="")

def remove_stopwords(text):
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(text)
    filtered_text = [word for word in word_tokens if word.lower() not in stop_words]
    return ' '.join(filtered_text)

def genaiModel(file, query):
    generation_config = {
    "temperature": 0.9,
    "top_p": 1,
    "top_k": 1,
    "max_output_tokens": 2048,
    }

    safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    ]
    model = genai.GenerativeModel(model_name="gemini-1.0-pro",
                                generation_config=generation_config,
                                safety_settings=safety_settings)
    root, extension = os.path.splitext(file)
    file = root
    text = open('output/' + file + '.txt', "r+", encoding="utf8", errors="ignore")

    que= query
    prompt_parts = [
    text.read(),
    "Answer the questions only if they are related to given paper. Only Extract sentences from the given text do not add anything of your own.",
    query,
    ]

    response = model.generate_content(prompt_parts)
    return response.text


def genaiModel2(file, query):
    generation_config = {
  "temperature": 0.4,
  "top_p": 1,
  "top_k": 32,
  "max_output_tokens": 4096,
}

    safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
    },
    ]

    model = genai.GenerativeModel(model_name="gemini-1.0-pro-vision-latest",
                                generation_config=generation_config,
                                safety_settings=safety_settings)

    # Validate that an image is present
    image_parts = [
    {
        "mime_type": "image/jpeg",
        "data": Path("/Users/varadunhale/Desktop/New_paperInsight/data/"+file).read_bytes()
    },
    ]
    prompt_parts = [
    query,image_parts[0]
    ]
    response = model.generate_content(prompt_parts)
    return response.text
def upload_files(request):
    print('upload file')
    output = ""
    file = None
    context = {
        'file' : "Not uploaded",
        'output' : output
    }
    if request.method == 'POST':
        if 'upload' in request.POST:
            for file in request.FILES.getlist('document'):
                file_name = re.sub(r"(\s)|(\")|(')|(&)", "_", str(file.name))
                file_name = re.sub(r'%22','_',f'{file_name}')
                file.name = file_name
                fs = FileSystemStorage()
                
                path = "data/"+file.name
                extension = pathlib.Path(path).suffix   

                print(extension)           
                if extension == ".pdf":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)
                    convertapi.convert('txt', {
                        'File': path
                    }, from_format = 'pdf').save_files('output')
                    context = {
                        'file' : file,
                        'output' : ""
                    }
                elif extension == ".pptx" or extension ==  ".ppt":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)
                    convertapi.convert('pdf', {
                        'File': path
                    }, from_format = 'pptx').save_files('pdf')
                    print("fileName:" , file.name)
                    convertapi.convert('txt', {
                        'File': 'pdf/'+ 'Host-afe.pdf'
                    }, from_format = 'pdf').save_files('output')
                    context = {
                        'file' : file,
                        'output' : ""
                    }
                elif extension == ".docx" :
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    filename = path
                    document = docx.Document(filename)
                    text_file_path = 'output\\' + filename[5:-5] + '.txt'
                    print(text_file_path + "!!!")

                    with open(text_file_path, "w", encoding="utf-8") as f:

                        f.write("**Text:**\n")
                        for paragraph in document.paragraphs:
                            f.write(paragraph.text.strip() + "\n")

                        f.write("\n**Tables:**\n")
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    f.write(cell.text.strip() + "\t")
                                f.write("\n")
                    context = {
                        'file' : file,
                        'output' : ""
                    }
                elif extension == ".tex":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    context = {
                        'file' : file,
                    }
                elif extension == ".jpeg":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    context = {
                        'file' : file,
                    }
                elif extension == ".png":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    context = {
                        'file' : file,
                    }
                return render(request, 'upload.html', context)

        else:
            message = request.POST.get('message')
            file = request.POST.get('file')
            output = genaiModel(file, message)
            context = {
                'file' : file,
                'output' : output,
                'que' : message,
            }
            
    return render(request, 'upload.html', context)


def upload_files2(request):
    print('upload file')
    output = ""
    file = None
    context = {
        'file' : "Not uploaded",
        'output' : output
    }
    if request.method == 'POST':
        if 'upload' in request.POST:
            for file in request.FILES.getlist('document'):
                file_name = re.sub(r"(\s)|(\")|(')|(&)", "_", str(file.name))
                file_name = re.sub(r'%22','_',f'{file_name}')
                file.name = file_name
                fs = FileSystemStorage()
                
                path = "data/"+file.name
                extension = pathlib.Path(path).suffix   

                print(extension)           

                if extension == ".jpeg":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    context = {
                        'file' : file,
                    }
                elif extension == ".png":
                    if fs.exists(file.name):
                        fs.delete(file.name)
                    fs.save(file.name, file)

                    context = {
                        'file' : file,
                    }
                return render(request, 'upload.html', context)

        else:
            message = request.POST.get('message')
            file = request.POST.get('file')
            output = genaiModel2(file, message)
            context = {
                'file' : file,
                'output' : output,
                'que' : message,
            }

                
    return render(request, 'uploadvid.html', context)